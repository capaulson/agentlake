"""DOCX file adapter using python-docx."""

from __future__ import annotations

from pathlib import Path

from agentlake.adapters.base import ExtractedContent, StructureMarker, TextBlock


class DocxAdapter:
    """Extracts text from Microsoft Word ``.docx`` files.

    Paragraphs and tables are emitted as separate :class:`TextBlock` entries
    with locators ``paragraph:{n}`` and ``table:{n}``.
    """

    supported_extensions: list[str] = [".docx"]
    supported_mimetypes: list[str] = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]

    def can_handle(self, filename: str, content_type: str) -> bool:
        return (
            Path(filename).suffix.lower() in self.supported_extensions
            or content_type in self.supported_mimetypes
        )

    def extract(self, file_bytes: bytes, filename: str) -> ExtractedContent:
        """Extract paragraphs and tables from a DOCX file.

        Args:
            file_bytes: Raw DOCX bytes.
            filename: Original filename.

        Returns:
            Extracted content with paragraphs and table blocks.
        """
        import io

        import docx

        doc = docx.Document(io.BytesIO(file_bytes))

        metadata: dict = {"filename": filename}
        core = doc.core_properties
        if core.title:
            metadata["title"] = core.title
        if core.author:
            metadata["author"] = core.author
        if core.subject:
            metadata["subject"] = core.subject

        blocks: list[TextBlock] = []
        structure: list[StructureMarker] = []
        position = 0
        para_count = 0
        table_count = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            para_count += 1
            style_name = (para.style.name or "").lower() if para.style else ""

            if "heading" in style_name:
                # Attempt to extract heading level
                level = 1
                for ch in style_name:
                    if ch.isdigit():
                        level = int(ch)
                        break
                blocks.append(
                    TextBlock(
                        content=text,
                        block_type="heading",
                        position=position,
                        source_locator=f"paragraph:{para_count}",
                        metadata={"level": level, "style": para.style.name},
                    )
                )
                structure.append(
                    StructureMarker(
                        marker_type="section_start",
                        position=position,
                        metadata={"level": level, "title": text},
                    )
                )
            else:
                blocks.append(
                    TextBlock(
                        content=text,
                        block_type="paragraph",
                        position=position,
                        source_locator=f"paragraph:{para_count}",
                        metadata={"style": para.style.name if para.style else "Normal"},
                    )
                )
            position += 1

        for table in doc.tables:
            table_count += 1
            rows_text: list[str] = []

            structure.append(
                StructureMarker(
                    marker_type="table_start",
                    position=position,
                    metadata={"table_index": table_count},
                )
            )

            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))

            table_content = "\n".join(rows_text)
            if table_content.strip():
                blocks.append(
                    TextBlock(
                        content=table_content,
                        block_type="table",
                        position=position,
                        source_locator=f"table:{table_count}",
                        metadata={"row_count": len(table.rows), "col_count": len(table.columns)},
                    )
                )

            structure.append(
                StructureMarker(
                    marker_type="table_end",
                    position=position,
                    metadata={"table_index": table_count},
                )
            )
            position += 1

        return ExtractedContent(
            text_blocks=blocks,
            metadata=metadata,
            structure=structure,
        )
